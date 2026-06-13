from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document


PROJECT_DIR = Path(__file__).resolve().parents[1]
AI_DIR = Path(__file__).resolve().parent
SOURCE_CONFIG = AI_DIR / "knowledge_sources.json"
OUTPUT_FILE = AI_DIR / "data" / "chunks.json"

CHUNK_SIZE = 1_000
CHUNK_OVERLAP = 150


def normalise_text(text: str) -> str:
    """Remove repeated whitespace while preserving paragraph boundaries."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
    return "\n".join(line for line in lines if line)


def load_docx(path: Path) -> list[dict[str, str]]:
    document = Document(path)
    sections: list[dict[str, str]] = []
    current_heading = path.stem
    buffer: list[str] = []

    def flush() -> None:
        nonlocal buffer
        text = normalise_text("\n".join(buffer))
        if text:
            sections.append({"heading": current_heading, "text": text})
        buffer = []

    for paragraph in document.paragraphs:
        text = normalise_text(paragraph.text)
        if not text:
            continue
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            flush()
            current_heading = text
        else:
            buffer.append(text)

    # Include table content because important methodology is stored in tables.
    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            values = [normalise_text(cell.text) for cell in row.cells]
            rows.append(" | ".join(values))
        table_text = normalise_text("\n".join(rows))
        if table_text:
            buffer.append(f"Table {table_index}\n{table_text}")

    flush()
    return sections


def load_text_file(path: Path) -> list[dict[str, str]]:
    text = path.read_text(encoding="utf-8-sig")
    text = normalise_text(text)
    return [{"heading": path.stem, "text": text}]


def split_text(text: str, size: int, overlap: int) -> list[str]:
    """Character splitter with overlap and sentence-aware boundaries."""
    if overlap >= size:
        raise ValueError("CHUNK_OVERLAP must be smaller than CHUNK_SIZE.")

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + size, len(text))
        if end < len(text):
            boundary = max(
                text.rfind("\n", start, end),
                text.rfind(". ", start, end),
                text.rfind("。", start, end),
                text.rfind("; ", start, end),
            )
            if boundary > start + size // 2:
                end = boundary + 1

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = max(0, end - overlap)

    return chunks


def main() -> None:
    configured_sources = json.loads(SOURCE_CONFIG.read_text(encoding="utf-8"))
    records: list[dict[str, object]] = []

    for configured_path in configured_sources:
        source_path = (AI_DIR / configured_path).resolve()
        if not source_path.exists():
            raise FileNotFoundError(f"Knowledge source not found: {source_path}")

        if source_path.suffix.lower() == ".docx":
            sections = load_docx(source_path)
        elif source_path.suffix.lower() in {".md", ".sql", ".txt"}:
            sections = load_text_file(source_path)
        else:
            raise ValueError(f"Unsupported source type: {source_path.suffix}")

        for section_index, section in enumerate(sections):
            for chunk_index, chunk in enumerate(
                split_text(section["text"], CHUNK_SIZE, CHUNK_OVERLAP)
            ):
                records.append(
                    {
                        "chunk_id": f"chunk-{len(records) + 1:04d}",
                        "source": str(source_path.relative_to(PROJECT_DIR)),
                        "file_name": source_path.name,
                        "heading": section["heading"],
                        "section_index": section_index,
                        "chunk_index": chunk_index,
                        "text": chunk,
                        "character_count": len(chunk),
                    }
                )

    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_FILE.write_text(
        json.dumps(records, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    print(f"Sources loaded: {len(configured_sources)}")
    print(f"Chunks created: {len(records)}")
    print(f"Output: {OUTPUT_FILE}")
    print("\nSample chunk:")
    print(json.dumps(records[0], ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
