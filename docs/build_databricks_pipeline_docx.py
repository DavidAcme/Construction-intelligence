from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
OUTPUT_PATH = BASE_DIR / "docs" / "Adbri_Excel_to_Databricks_SQL_Pipeline_CN_EN.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
GRAY_FILL = "F2F4F7"


def style_run(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_para(doc, text, label=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if label and text.startswith(label):
        r1 = p.add_run(label)
        style_run(r1, bold=True)
        r2 = p.add_run(text[len(label):])
        style_run(r2)
    else:
        r = p.add_run(text)
        style_run(r)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    r = p.add_run(text)
    style_run(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True, color=BLUE if level <= 2 else DARK_BLUE)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    for line in text.strip("\n").split("\n"):
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(9)
        p.add_run("\n")


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        style_run(r)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    style_run(r, size=10, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_fill(cell, GRAY_FILL)
        set_cell_text(cell, header, bold=True)
        cell.width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].width = Inches(widths[i])
    doc.add_paragraph()


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color in [("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, DARK_BLUE)]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Adbri Excel to Databricks SQL Pipeline")
    style_run(r, size=22, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("四个 ABS Excel 如何清洗成 CSV、上传 Databricks、用 SQL 建模、连接 Power BI")
    style_run(r, size=11, color=RGBColor(85, 85, 85))

    add_heading(doc, "1. Executive Summary | 总体思路", 1)
    add_para(doc, "English: The four Excel files are ABS time-series workbooks. They are not clean database tables. Each workbook contains metadata rows and many time-series columns. I first reshaped them into one tidy long CSV, then designed a Databricks medallion pipeline: Bronze for raw data, Silver for typed and cleaned data, and Gold for Power BI reporting tables.", "English:")
    add_para(doc, "中文：这四个 Excel 是 ABS 的时间序列 workbook，不是普通数据库表。每个文件都有 metadata 行和很多横向展开的 series 列。我先把它们清洗成一个 tidy long CSV，然后设计 Databricks medallion pipeline：Bronze 原始层、Silver 清洗类型层、Gold Power BI 报表层。", "中文：")
    add_code(doc, "4 Excel files -> Python reshape -> Combined CSV -> Databricks Volume -> Bronze -> Silver -> Gold -> Power BI")

    add_heading(doc, "2. Source Excel Structure | 原始 Excel 结构", 1)
    add_para(doc, "English: Each source workbook has three sheets: Index, Data1, and Enquiries. The real time-series data is in Data1.", "English:")
    add_para(doc, "中文：每个源文件都有三个 sheet：Index、Data1、Enquiries。真正用于分析的数据在 Data1。", "中文：")
    add_table(
        doc,
        ["Rows in Data1", "Meaning | 含义"],
        [
            ["Row 1", "Series description | 指标描述"],
            ["Row 2", "Unit | 单位"],
            ["Row 3", "Series Type | Original / Seasonally Adjusted 等"],
            ["Row 4", "Data Type | FLOW 或 INDEX"],
            ["Row 5", "Frequency | Month 或 Quarter"],
            ["Rows 6-9", "Collection month, start/end dates, observation count | 元数据"],
            ["Row 10", "Series ID | ABS 指标 ID"],
            ["Row 11 onwards", "Date and values | 日期和值"],
        ],
        [1.5, 4.8],
    )

    add_heading(doc, "3. Why Convert Wide Excel to Long CSV | 为什么要转成长表 CSV", 1)
    add_para(doc, "English: The original Data1 sheet is a wide time-series table. Each series is a separate column. Databricks and Power BI work better with a long fact table, where each row is one observation for one series and one period.", "English:")
    add_para(doc, "中文：原始 Data1 是宽表，每个 series 是一列。Databricks 和 Power BI 更适合长表，也就是每一行代表一个 series 在一个时间点的一个观测值。", "中文：")
    add_code(doc, """Original wide format:
Date        Series A   Series B   Series C
1983-07     535        625        111
1983-08     628        720        186

Tidy long format:
period_start   series_key                    value
1983-07-01     table04|A418704W              535
1983-07-01     table04|A418757A              625
1983-07-01     table04|A421567W              111""")

    add_heading(doc, "4. Python Extraction Logic | Python 清洗逻辑", 1)
    add_para(doc, "English: The Python script reads the Data1 sheet from each workbook, extracts series metadata from rows 1 to 10, then loops through each date row and each series column to create a long-format record.", "English:")
    add_para(doc, "中文：Python 脚本读取每个 workbook 的 Data1 sheet，把第 1 到第 10 行解析成 series metadata，然后从第 11 行开始，按日期行和 series 列循环，生成长表记录。", "中文：")
    add_bullets(doc, [
        "Input files: four ABS Excel workbooks from the desktop source folder.",
        "Sheet used: Data1.",
        "Metadata parsed: series description, unit, series type, data type, frequency, series start/end, no. obs, series_id.",
        "Business attributes parsed: metric, category_1 to category_5, region, city.",
        "Output: one combined CSV plus four individual CSVs.",
    ])
    add_para(doc, "中文重点：脚本不仅提取 value，还保留 series metadata，并解析 category、region、city，这样 Power BI 里才有业务维度可以筛选。")
    add_code(doc, """Important output columns:
source_file, dataset, series_key, subject, abs_table,
period_start, year, quarter, month, value,
series_id, series_description, unit, frequency,
metric, category_1, category_2, category_3, region, city""")

    add_heading(doc, "5. Output CSV Files | 输出结果", 1)
    add_table(
        doc,
        ["CSV", "Rows"],
        [
            ["building_approvals_sa_table04.csv", "7,182"],
            ["building_approvals_sa_table15.csv", "9,382"],
            ["ppi_construction_output_table17.csv", "4,980"],
            ["ppi_house_construction_input_table18.csv", "22,182"],
            ["adbri_abs_ppi_combined_long.csv", "43,726"],
        ],
        [4.7, 1.2],
    )
    add_para(doc, "English: The combined CSV is the file uploaded to Databricks. It contains all observations in one consistent schema.", "English:")
    add_para(doc, "中文：上传到 Databricks 的是 combined CSV，它把四个文件统一成同一个 schema。", "中文：")

    add_heading(doc, "6. Why Use series_key | 为什么要用 series_key", 1)
    add_para(doc, "English: I did not use series_id alone as the relationship key. ABS series_id can appear across different tables, so I created a safer composite key: dataset + '|' + series_id.", "English:")
    add_para(doc, "中文：我没有直接用 series_id 做关系键，因为 ABS 的同一个 series_id 可能跨表出现。更安全的方式是用 dataset + '|' + series_id 生成 series_key。", "中文：")
    add_code(doc, """series_key = dataset + "|" + series_id

Example:
building_approvals_sa_table04|A418704W
ppi_house_construction_input_table18|A2390561L""")
    add_para(doc, "Interview line: I created a composite series_key to avoid duplicate-key issues across ABS tables.", "Interview line:")

    add_heading(doc, "7. Uploading to Databricks | 上传到 Databricks", 1)
    add_para(doc, "English: In Databricks, I would upload the combined CSV to a Unity Catalog Volume. Volumes are useful because they provide governed file storage under a catalog and schema.", "English:")
    add_para(doc, "中文：在 Databricks 里，我会把 combined CSV 上传到 Unity Catalog Volume。Volume 的好处是文件存储也在 catalog/schema 管理下，更适合治理。", "中文：")
    add_bullets(doc, [
        "Create catalog: adbri_bi.",
        "Create schemas: raw, silver, gold.",
        "Create volume: /Volumes/adbri_bi/raw/adbri_raw/.",
        "Upload file: adbri_abs_ppi_combined_long.csv.",
    ])
    add_code(doc, """/Volumes/adbri_bi/raw/adbri_raw/adbri_abs_ppi_combined_long.csv""")

    add_heading(doc, "8. Bronze Layer SQL | Bronze 原始层", 1)
    add_para(doc, "English: Bronze keeps the raw uploaded CSV as a table with minimal transformation. This gives us a reproducible starting point.", "English:")
    add_para(doc, "中文：Bronze 层尽量保留 CSV 原貌，只做最少转换。这样后续所有清洗都有可追溯的起点。", "中文：")
    add_code(doc, """CREATE CATALOG IF NOT EXISTS adbri_bi;
CREATE SCHEMA IF NOT EXISTS adbri_bi.raw;
CREATE SCHEMA IF NOT EXISTS adbri_bi.silver;
CREATE SCHEMA IF NOT EXISTS adbri_bi.gold;

CREATE OR REPLACE TABLE raw.abs_ppi_long_raw
USING CSV
OPTIONS (
  path '/Volumes/adbri_bi/raw/adbri_raw/adbri_abs_ppi_combined_long.csv',
  header 'true',
  inferSchema 'true',
  encoding 'UTF-8'
);""")

    add_heading(doc, "9. Silver Layer SQL | Silver 清洗层", 1)
    add_para(doc, "English: Silver converts raw string-like fields into typed analytical columns. Dates become DATE, year/quarter/month become integers, and value becomes DOUBLE. This layer also keeps all useful metadata.", "English:")
    add_para(doc, "中文：Silver 层把原始字段转换成分析可用的类型。日期转 DATE，year/quarter/month 转整数，value 转 DOUBLE，同时保留 series metadata。", "中文：")
    add_code(doc, """CREATE OR REPLACE TABLE silver.fact_abs_ppi_series AS
SELECT
  source_file,
  dataset,
  COALESCE(series_key, CONCAT(dataset, '|', series_id)) AS series_key,
  subject,
  abs_table,
  CAST(period_start AS DATE) AS period_start,
  CAST(year AS INT) AS year,
  CAST(quarter AS INT) AS quarter,
  CAST(month AS INT) AS month,
  CAST(value AS DOUBLE) AS value,
  series_id,
  series_description,
  unit,
  series_type,
  data_type,
  frequency,
  metric,
  category_1,
  category_2,
  category_3,
  region,
  city,
  current_timestamp() AS loaded_at
FROM raw.abs_ppi_long_raw
WHERE period_start IS NOT NULL
  AND value IS NOT NULL;""")

    add_heading(doc, "10. Gold Layer SQL | Gold 报表层", 1)
    add_para(doc, "English: Gold is designed for Power BI. I created a star schema with one fact table and two dimensions: dim_date and dim_series.", "English:")
    add_para(doc, "中文：Gold 层是给 Power BI 用的。我设计了一个星型模型：一个事实表，两个维度表 dim_date 和 dim_series。", "中文：")
    add_code(doc, """CREATE OR REPLACE TABLE gold.dim_date AS
SELECT DISTINCT
  period_start AS date_key,
  year,
  quarter,
  month,
  date_format(period_start, 'yyyy-MM') AS year_month,
  CONCAT(CAST(year AS STRING), '-Q', CAST(quarter AS STRING)) AS year_quarter
FROM silver.fact_abs_ppi_series;

CREATE OR REPLACE TABLE gold.dim_series AS
SELECT DISTINCT
  series_key,
  series_id,
  dataset,
  subject,
  abs_table,
  series_description,
  unit,
  series_type,
  data_type,
  frequency,
  metric,
  category_1,
  category_2,
  category_3,
  region,
  city
FROM silver.fact_abs_ppi_series;""")
    add_code(doc, """CREATE OR REPLACE TABLE gold.fact_market_indicators AS
SELECT
  period_start AS date_key,
  series_key,
  series_id,
  dataset,
  subject,
  frequency,
  region,
  city,
  unit,
  value,
  LAG(value, 1) OVER (PARTITION BY series_key ORDER BY period_start) AS prior_period_value,
  LAG(value, 4) OVER (PARTITION BY series_key ORDER BY period_start) AS value_4_periods_ago,
  CASE
    WHEN LAG(value, 1) OVER (PARTITION BY series_key ORDER BY period_start) = 0 THEN NULL
    ELSE value / LAG(value, 1) OVER (PARTITION BY series_key ORDER BY period_start) - 1
  END AS period_change_pct,
  CASE
    WHEN LAG(value, 4) OVER (PARTITION BY series_key ORDER BY period_start) = 0 THEN NULL
    ELSE value / LAG(value, 4) OVER (PARTITION BY series_key ORDER BY period_start) - 1
  END AS annual_change_pct
FROM silver.fact_abs_ppi_series;""")

    add_heading(doc, "11. Data Quality Checks | 数据质量检查", 1)
    add_para(doc, "English: I added quality checks to make the pipeline auditable: row counts, distinct series counts, date coverage, null checks, and duplicate grain checks.", "English:")
    add_para(doc, "中文：我加入了数据质量检查，让 pipeline 可审计，包括行数、唯一 series 数、日期范围、空值、重复粒度检查。", "中文：")
    add_code(doc, """-- Duplicate grain check
SELECT
  series_key,
  period_start,
  COUNT(*) AS duplicate_count
FROM silver.fact_abs_ppi_series
GROUP BY series_key, period_start
HAVING COUNT(*) > 1;

-- Null critical fields
SELECT
  SUM(CASE WHEN series_key IS NULL THEN 1 ELSE 0 END) AS null_series_key,
  SUM(CASE WHEN period_start IS NULL THEN 1 ELSE 0 END) AS null_period_start,
  SUM(CASE WHEN value IS NULL THEN 1 ELSE 0 END) AS null_value
FROM silver.fact_abs_ppi_series;""")
    add_para(doc, "Interview line: The fact table grain is one row per series_key per period_start, and I validated that with a duplicate-grain SQL check.", "Interview line:")

    add_heading(doc, "12. Connecting Power BI | Power BI 连接方式", 1)
    add_para(doc, "English: Power BI should connect to the Gold tables, not raw Excel. This keeps the report stable, governed, and easier to maintain.", "English:")
    add_para(doc, "中文：Power BI 应该连接 Gold 表，而不是直接连原始 Excel。这样报表更稳定、更易维护，也更符合数据治理。", "中文：")
    add_bullets(doc, [
        "Power BI Desktop -> Get Data -> Azure Databricks.",
        "Enter Databricks Server Hostname and HTTP Path.",
        "Choose Import for portfolio demo, or DirectQuery for near-real-time production use.",
        "Load gold.fact_market_indicators, gold.dim_date, and gold.dim_series.",
    ])
    add_table(
        doc,
        ["Relationship", "Cardinality"],
        [
            ["fact_market_indicators[date_key] -> dim_date[date_key]", "Many-to-one"],
            ["fact_market_indicators[series_key] -> dim_series[series_key]", "Many-to-one"],
        ],
        [4.7, 1.4],
    )

    add_heading(doc, "13. Full English Interview Script | 英文完整话术", 1)
    add_para(doc, "The original ABS Excel files were time-series workbooks, not clean database tables. The actual observations were stored in the Data1 sheet, with the first ten rows containing metadata such as unit, frequency, series type, data type, date range, and ABS series ID. From row eleven onwards, the data was stored in a wide format, where each series was a separate column.")
    add_para(doc, "I used Python to read each workbook, extract the metadata, and reshape the data from wide format into a tidy long CSV. Each output row represents one observation for one series and one period. I also created a composite series_key using dataset and series_id, because the same ABS series_id can appear across different tables.")
    add_para(doc, "In Databricks, I would upload the combined CSV into a Unity Catalog Volume and create a Bronze raw table from the file. Then I would build a Silver table where dates and numeric values are properly typed and the metadata is preserved. Finally, I would publish Gold tables for Power BI: a market indicators fact table, a date dimension, and a series dimension.")
    add_para(doc, "The Gold layer supports a Power BI star schema. The fact table grain is one row per series_key per period_start. The date dimension connects through date_key and the series dimension connects through series_key. I also added SQL quality checks for row counts, null critical fields, duplicate grain, and date coverage. This makes the report explainable, auditable, and easier to maintain.")

    add_heading(doc, "14. 中文完整话术", 1)
    add_para(doc, "这四个 ABS Excel 原始文件不是普通数据库表，而是 time-series workbook。真正的数据在 Data1 sheet。前 10 行是 metadata，例如单位、频率、series type、data type、起止日期、ABS series ID。从第 11 行开始才是日期和值，而且原始结构是宽表，每个 series 是一列。")
    add_para(doc, "我用 Python 读取每个 workbook，解析前 10 行 metadata，然后把宽表转换成长表。转换后，每一行代表一个 series 在一个 period 的一个观测值。我还创建了 composite key，也就是 series_key = dataset + series_id，因为 ABS 的同一个 series_id 有可能出现在不同表里，直接用 series_id 做关系不够安全。")
    add_para(doc, "如果放到 Databricks，我会先把 combined CSV 上传到 Unity Catalog Volume，然后用 SQL 创建 Bronze raw table。接着创建 Silver 表，把日期、数值、year、quarter、month 等字段转换成正确类型，并保留 metadata。最后创建 Gold 层给 Power BI 使用，包括 fact_market_indicators、dim_date 和 dim_series。")
    add_para(doc, "Gold 层支持 Power BI 星型模型。事实表粒度是 one row per series_key per period_start。date_key 连接日期维度，series_key 连接指标维度。我还加入了 SQL 数据质量检查，例如行数检查、关键字段空值检查、重复粒度检查和日期范围检查。这样报表不仅能展示图表，也能解释数据来源和质量。")

    add_heading(doc, "15. Common Questions | 常见问题回答", 1)
    add_table(
        doc,
        ["Question | 问题", "Answer | 回答"],
        [
            ["Why not connect Power BI directly to Excel?", "Databricks gives governed, reusable, typed Gold tables. Direct Excel is harder to maintain. | Databricks 可以提供治理过、可复用、类型清晰的 Gold 表；直接连 Excel 难维护。"],
            ["Why use Bronze/Silver/Gold?", "It separates raw ingestion, cleaning, and reporting layers. | 它把原始采集、清洗转换、报表服务分开。"],
            ["Why convert wide to long?", "Long format is the right fact-table shape for SQL, Databricks, and Power BI. | 长表更适合事实表、SQL 分析和 Power BI 建模。"],
            ["Why create series_key?", "To avoid duplicate-key problems when the same ABS series_id appears across datasets. | 避免同一个 ABS series_id 跨数据集重复导致关系错误。"],
            ["What quality checks did you add?", "Row counts, null checks, duplicate grain checks, and latest period checks. | 行数、空值、重复粒度、最新日期检查。"],
        ],
        [2.2, 4.1],
    )

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_doc()
