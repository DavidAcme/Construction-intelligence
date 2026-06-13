from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
ASSET_DIR = BASE_DIR / "docs" / "final_report_assets"
DELIVERY_DIR = Path(os.environ.get("ADBRI_DELIVERY_DIR", BASE_DIR / "docs"))

EN_OUT = BASE_DIR / "docs" / "Adbri_Market_Demand_BI_Research_Report_EN.docx"
CN_OUT = BASE_DIR / "docs" / "Adbri_Market_Demand_BI_Research_Report_CN.docx"

BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(33, 33, 33)
MUTED = RGBColor(90, 90, 90)
HEADER_FILL = "EAF1F8"


def set_run(run, size=11, bold=False, italic=False, color=None, chinese=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun" if chinese else "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def setup_doc(chinese=False) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    east = "SimSun" if chinese else "Times New Roman"
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    normal = doc.styles["Normal"]
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)
    for name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)]:
        s = doc.styles[name]
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = BLUE if name != "Heading 3" else DARK
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after = Pt(6)
    return doc


def title(doc, main, subtitle, chinese=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(main)
    set_run(r, 21, True, color=BLUE, chinese=chinese)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_run(r, 11, color=MUTED, chinese=chinese)


def h(doc, text, level=1, chinese=False):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_run(r, {1: 16, 2: 13, 3: 12}.get(level, 11), True, color=BLUE if level < 3 else DARK, chinese=chinese)


def p(doc, text, chinese=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    r = para.add_run(text)
    set_run(r, chinese=chinese)


def bullets(doc, items, chinese=False):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(4)
        r = para.add_run(item)
        set_run(r, chinese=chinese)


def shade(cell, fill=HEADER_FILL):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def table(doc, headers, rows, widths, chinese=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    for i, text in enumerate(headers):
        cell = t.rows[0].cells[i]
        shade(cell)
        cell.width = Inches(widths[i])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        r = cell.paragraphs[0].add_run(text)
        set_run(r, 10, True, chinese=chinese)
    for row in rows:
        cells = t.add_row().cells
        for i, text in enumerate(row):
            cells[i].width = Inches(widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            r = cells[i].paragraphs[0].add_run(str(text))
            set_run(r, 10, chinese=chinese)
    doc.add_paragraph()


def figure(doc, image_name, caption, chinese=False):
    path = ASSET_DIR / image_name
    doc.add_picture(str(path), width=Inches(6.35))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run(r, 9, italic=True, color=MUTED, chinese=chinese)


def build_en():
    doc = setup_doc()
    title(
        doc,
        "A Power BI-Based Market Demand and Cost Pressure Study for Adbri",
        "An applied business intelligence report using ABS building approvals and producer price index data",
    )

    h(doc, "Abstract")
    p(doc, "This report presents a Power BI-based business intelligence study designed to evaluate construction market demand and cost pressure relevant to Adbri, a building materials company operating in South Australia. Four Australian Bureau of Statistics Excel time-series workbooks were transformed into an analytical dataset and modelled through a Databricks-style Bronze, Silver, and Gold architecture. Power BI was then used as the analytical and presentation layer to construct an executive market pulse, a demand analysis page, a cost pressure page, and a data quality and lineage page. The resulting report indicates that South Australia building approvals recovered in 2025 after softer conditions in 2022-2023, while construction PPI measures continued to show positive cost and price pressure into March 2026.")

    h(doc, "1. Introduction")
    p(doc, "The purpose of the report is not merely to demonstrate dashboard construction. The central question is whether Power BI can be used to derive a defensible market view for a building materials company. For Adbri, the relevant commercial question is whether construction activity is likely to support demand for materials such as cement, concrete, aggregates, lime, masonry, and related products, while also understanding whether construction cost pressure remains elevated.")
    p(doc, "The study therefore combines two categories of public market indicators: building approvals as a forward-looking demand proxy, and Producer Price Index measures as indicators of construction input and output price pressure.")

    h(doc, "2. Data and Methodology")
    p(doc, "The source data consists of four ABS Excel workbooks: two building approval tables for South Australia and two PPI tables covering construction output and house construction input costs. The original Excel files are not database-ready tables; each workbook stores series metadata in the first rows and observations in a wide time-series layout. A Python extraction process reshaped the workbooks into a long-format fact table suitable for SQL modelling and Power BI analysis.")
    table(doc, ["Stage", "Methodological Role"], [
        ["Excel source", "ABS building approvals and PPI time-series workbooks"],
        ["Python transformation", "Wide-to-long reshaping, metadata extraction, creation of series_key"],
        ["Databricks design", "Bronze raw table, Silver typed fact table, Gold star-schema tables"],
        ["Power BI analysis", "Interactive visual analysis, DAX measures, slicers, KPI cards, demand and PPI interpretation"],
    ], [2.0, 4.3])
    p(doc, "The Power BI model uses a star schema. The fact table grain is one observation per series_key per period_start. The series dimension stores dataset, subject, frequency, category, region, and city attributes; the date dimension stores year, quarter, month, and period labels. This model supports filtered analysis across demand, geography, frequency, and cost indicator type.")

    h(doc, "3. Power BI Analytical Framework")
    p(doc, "Power BI was used as the analytical layer, not simply as a charting tool. The report combines DAX measures, relationship-driven filter propagation, slicer-controlled context, and page-specific analytical designs. The following figures show the four pages used to derive and communicate the market findings.")
    figure(doc, "powerbi_page_1_executive.png", "Figure 1. Executive Market Pulse page showing overall demand and cost-pressure indicators.")
    p(doc, "The Executive Market Pulse page consolidates latest market value, annual change, period change, and long-term trend analysis. This page is designed for a senior stakeholder who needs a fast view of whether the market is expanding, softening, or facing price pressure. The use of slicers for subject, frequency, and dataset allows the same model to support multiple analytical lenses.")
    figure(doc, "powerbi_page_2_demand.png", "Figure 2. Building Approvals Demand page showing houses and dwellings excluding houses.")
    p(doc, "The Building Approvals Demand page isolates South Australia building approvals. By separating houses from dwellings excluding houses, the report identifies the composition of construction demand rather than presenting only a total. This page is the strongest evidence for market demand interpretation because building approvals are a leading indicator of future materials demand.")
    figure(doc, "powerbi_page_3_cost.png", "Figure 3. Construction Cost Pressure page showing Adelaide input PPI and South Australia output PPI.")
    p(doc, "The Construction Cost Pressure page separates two PPI series with different geographic grains: Adelaide house construction input PPI and South Australia construction output PPI. This design decision avoids mixing city-level and state-level measures in a single visual and demonstrates disciplined Power BI modelling.")
    figure(doc, "powerbi_page_4_quality.png", "Figure 4. Data Quality & Lineage page documenting row counts, series counts, date coverage, and model grain.")
    p(doc, "The Data Quality & Lineage page makes the analytical report auditable. It shows fact rows, unique series count, date periods, source dataset counts, data pipeline steps, and the model grain. This gives the viewer confidence that the dashboard is supported by a traceable data process.")

    h(doc, "4. Results")
    h(doc, "4.1 Demand Conditions", 2)
    p(doc, "The Power BI demand page indicates that South Australia building approvals recovered in 2025 after weaker annual totals in 2022 and 2023. Annual approvals for houses and dwellings excluding houses were approximately 24,474 in 2024 and 29,623 in 2025, indicating a material recovery in approval activity. At the latest monthly observation in March 2026, total approvals for the selected building-type categories were approximately 2,594, compared with 2,388 in March 2025, an increase of approximately 8.6%.")
    p(doc, "The composition of demand is uneven. Houses were slightly lower year-on-year in March 2026, while dwellings excluding houses were materially higher. This suggests that the demand signal is not uniformly house-led; multi-unit or non-house dwelling categories appear to be contributing more strongly to the latest improvement.")
    h(doc, "4.2 Cost and Price Pressure", 2)
    p(doc, "The PPI page shows that cost and output price pressure remained positive into March 2026. Adelaide house construction input PPI reached 164.7 in March 2026, approximately 3.7% higher than March 2025. South Australia building construction output PPI reached 148.9, approximately 3.8% higher than March 2025. These results indicate that the market is not only showing demand recovery but also continued price pressure.")

    h(doc, "5. Discussion")
    p(doc, "For Adbri, the findings imply a cautiously positive demand environment. The recovery in 2025 approvals supports the view that construction-related materials demand may strengthen relative to the weaker 2022-2023 period. However, the demand mix matters. If growth is concentrated in dwellings excluding houses rather than detached houses, product mix and customer segment implications should be examined more closely.")
    p(doc, "The PPI evidence suggests that pricing and cost conditions remain important. A positive demand signal combined with persistent input and output price pressure means commercial teams should monitor both volume opportunity and margin risk. This is exactly where a Power BI report adds business value: it allows stakeholders to move from broad market questions to filtered evidence by subject, frequency, category, and geography.")

    h(doc, "6. Conclusion")
    p(doc, "The report demonstrates that a Power BI-driven BI workflow can convert complex ABS Excel workbooks into a business-ready market intelligence asset. The analysis suggests that South Australia construction demand improved in 2025 and remained supported into early 2026, while PPI indicators show continued construction cost and output price pressure. For an Adbri-style building materials business, the combined evidence points to recovering market demand with ongoing pricing pressure, requiring close monitoring of both volume and margin conditions.")

    h(doc, "Appendix: Interview Explanation")
    p(doc, "In an interview, I would explain that I used Power BI not only to build visuals, but to structure a market analysis. I created a star schema, wrote DAX measures, built slicers and KPI cards, separated visuals by business question, and included a data quality page. The report shows demand recovery, different demand composition between houses and other dwellings, and continued construction PPI pressure.")
    doc.save(EN_OUT)


def build_cn():
    doc = setup_doc(chinese=True)
    title(
        doc,
        "基于 Power BI 的 Adbri 市场需求与成本压力研究报告",
        "使用 ABS 建筑审批与 Producer Price Index 数据的应用型商业智能分析",
        chinese=True,
    )

    h(doc, "摘要", chinese=True)
    p(doc, "本报告展示了一个面向 Adbri 这类建筑材料公司的 Power BI 商业智能研究。研究目标是判断 South Australia 建筑市场需求是否恢复，以及 Adelaide 和 South Australia 建筑成本压力是否仍然存在。项目将四个 ABS Excel 时间序列 workbook 清洗成统一分析数据集，并按照 Databricks Bronze、Silver、Gold 的思路建模，最后在 Power BI 中建立管理层总览、需求分析、成本压力分析、数据质量与血缘四个页面。分析结果显示，South Australia 建筑审批在 2025 年较 2022-2023 年低迷阶段出现恢复，同时 PPI 指标到 2026 年 3 月仍显示正向成本和价格压力。", chinese=True)

    h(doc, "1. 引言", chinese=True)
    p(doc, "本报告的目的不是简单展示 Power BI 图表，而是回答一个商业问题：我们能否用 Power BI 从公开建筑市场数据中得出对 Adbri 有意义的市场需求判断。对建筑材料公司而言，关键问题包括未来建筑活动是否支持材料需求，以及建筑投入成本和产出价格是否仍有压力。", chinese=True)
    p(doc, "因此，本研究结合两类指标：Building Approvals 作为未来建筑需求的前置信号，Producer Price Index 作为建筑投入成本和产出价格压力的指标。", chinese=True)

    h(doc, "2. 数据与方法", chinese=True)
    p(doc, "数据来自四个 ABS Excel workbook，包括两个 South Australia 建筑审批表，以及两个 PPI 表。原始 Excel 并不是数据库可直接使用的结构；每个 workbook 前几行存放 series metadata，后续数据以宽表形式展开。项目使用 Python 将宽表转换成长表，并生成适合 SQL 和 Power BI 分析的事实表。", chinese=True)
    table(doc, ["阶段", "方法作用"], [
        ["Excel 源数据", "ABS 建筑审批和 PPI 时间序列 workbook"],
        ["Python 转换", "宽表转长表、metadata 提取、series_key 生成"],
        ["Databricks 设计", "Bronze 原始层、Silver 清洗层、Gold 星型模型层"],
        ["Power BI 分析", "交互式分析、DAX 度量、slicers、KPI cards、需求和 PPI 解读"],
    ], [2.0, 4.3], chinese=True)
    p(doc, "Power BI 模型采用星型模型。事实表粒度是 one observation per series_key per period_start。series 维度保存 dataset、subject、frequency、category、region、city 等属性；date 维度保存 year、quarter、month 等时间属性。该模型支持按需求、地理、频率和指标类型进行筛选分析。", chinese=True)

    h(doc, "3. Power BI 分析框架", chinese=True)
    p(doc, "Power BI 在本项目中不是简单画图工具，而是分析层和展示层。报表结合了 DAX measures、关系模型筛选传播、slicer 控制的 filter context，以及按业务问题设计的页面。以下四张图展示了我们如何通过 Power BI 得出和表达市场结论。", chinese=True)
    figure(doc, "powerbi_page_1_executive.png", "图 1. Executive Market Pulse 页面：展示整体需求与成本压力指标。", chinese=True)
    p(doc, "第一页将最新值、年度变化、周期变化和长期趋势整合到一个管理层视图中。subject、frequency、dataset slicers 让同一个模型可以从不同角度分析市场。这一页回答的是：市场总体是在增长、放缓，还是面临价格压力。", chinese=True)
    figure(doc, "powerbi_page_2_demand.png", "图 2. Building Approvals Demand 页面：展示 Houses 与 Dwellings excluding houses。", chinese=True)
    p(doc, "第二页专门分析 South Australia 建筑审批需求。通过拆分 Houses 和 Dwellings excluding houses，报表不仅看总量，也看需求结构。这一页是判断 Adbri 市场需求最核心的页面，因为建筑审批是未来建筑材料需求的前置信号。", chinese=True)
    figure(doc, "powerbi_page_3_cost.png", "图 3. Construction Cost Pressure 页面：展示 Adelaide input PPI 与 South Australia output PPI。", chinese=True)
    p(doc, "第三页分析成本压力。Adelaide input PPI 是 city-level，South Australia output PPI 是 state-level，所以报表把它们拆成两个 visual，而不是强行放在同一个图里。这说明分析过程中考虑了数据粒度，而不是盲目堆图。", chinese=True)
    figure(doc, "powerbi_page_4_quality.png", "图 4. Data Quality & Lineage 页面：展示行数、series 数、日期覆盖和模型粒度。", chinese=True)
    p(doc, "第四页展示数据质量和血缘。它包含事实表行数、唯一 series 数、日期周期、按 dataset 的 series count、数据 pipeline 和模型粒度。这使报表不仅能展示结果，也能解释数据来源和可信度。", chinese=True)

    h(doc, "4. 研究结果", chinese=True)
    h(doc, "4.1 市场需求", 2, chinese=True)
    p(doc, "Power BI 需求页面显示，South Australia 建筑审批在 2025 年较 2022-2023 年低迷阶段出现恢复。Houses 与 Dwellings excluding houses 两类年度审批合计在 2024 年约为 24,474，在 2025 年约为 29,623，说明审批活动明显回升。最新月度数据为 2026 年 3 月，所选建筑类型审批合计约 2,594，而 2025 年 3 月约为 2,388，同比增长约 8.6%。", chinese=True)
    p(doc, "需求结构并不完全由 houses 推动。2026 年 3 月 houses 同比小幅下降约 1.8%，而 dwellings excluding houses 同比增长约 33.4%。这说明最新需求改善更多来自非独立屋或多单元住宅相关类别，而不是单一 houses 类别。", chinese=True)
    h(doc, "4.2 成本与价格压力", 2, chinese=True)
    p(doc, "PPI 页面显示，到 2026 年 3 月，建筑成本和产出价格压力仍然存在。Adelaide house construction input PPI 达到 164.7，较 2025 年 3 月约上升 3.7%；South Australia building construction output PPI 达到 148.9，较 2025 年 3 月约上升 3.8%。这说明市场不仅有需求恢复信号，也存在持续价格压力。", chinese=True)

    h(doc, "5. 讨论", chinese=True)
    p(doc, "对于 Adbri 来说，研究结果说明市场需求环境谨慎偏正面。2025 年建筑审批恢复，意味着相比 2022-2023 年的低迷阶段，未来建筑材料需求可能改善。但是需求结构很重要。如果增长更多来自 dwellings excluding houses，而不是 houses，那么产品组合、客户类型和销售渠道也需要进一步分析。", chinese=True)
    p(doc, "PPI 结果说明成本和价格压力仍然需要关注。需求恢复叠加投入成本和产出价格上升，意味着业务团队既要关注销量机会，也要关注毛利压力。Power BI 的价值就在这里：它可以把宏观市场问题转化成可筛选、可解释、可追溯的证据。", chinese=True)

    h(doc, "6. 结论", chinese=True)
    p(doc, "本报告证明，Power BI 驱动的 BI 工作流可以把复杂 ABS Excel workbook 转化为业务可用的市场洞察资产。分析显示，South Australia 建筑需求在 2025 年恢复，并延续到 2026 年初；同时 PPI 指标显示建筑成本和产出价格压力仍然存在。对 Adbri 这类建筑材料公司而言，综合证据指向：市场需求正在恢复，但价格和成本压力仍需持续监控。", chinese=True)

    h(doc, "附录：面试讲法", chinese=True)
    p(doc, "面试时可以这样讲：我不是单纯用 Power BI 做图，而是用 Power BI 组织市场分析。我建立星型模型，写 DAX measures，使用 slicers 和 KPI cards，按业务问题拆分页面，并加入数据质量页面。最后报表得出的结论是：South Australia 建筑审批在 2025 年恢复，需求结构中非 houses 类别增长更明显，同时 Adelaide 和 South Australia 的 PPI 仍显示成本和价格压力。", chinese=True)
    doc.save(CN_OUT)


def main():
    build_en()
    build_cn()
    for src, name in [
        (EN_OUT, "Adbri_Market_Demand_BI_Research_Report_EN.docx"),
        (CN_OUT, "Adbri_Market_Demand_BI_Research_Report_CN.docx"),
    ]:
        (DELIVERY_DIR / name).write_bytes(src.read_bytes())
        print(DELIVERY_DIR / name)


if __name__ == "__main__":
    main()
