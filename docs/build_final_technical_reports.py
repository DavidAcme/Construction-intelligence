from __future__ import annotations

import csv
import math
import os
from datetime import datetime
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
DATA_DIR = BASE_DIR / "powerbi" / "data"
ASSET_DIR = BASE_DIR / "docs" / "final_report_assets"
DOCS_DIR = BASE_DIR / "docs"
DELIVERY_DIR = Path(os.environ.get("ADBRI_DELIVERY_DIR", DOCS_DIR))

EN_DOC = DOCS_DIR / "Adbri_BI_Technical_Report_EN.docx"
CN_DOC = DOCS_DIR / "Adbri_BI_Technical_Report_CN.docx"

BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(31, 31, 31)
MUTED = RGBColor(89, 89, 89)
HEADER_FILL = "EAF1F8"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        r"C:\Windows\Fonts\times.ttf",
        r"C:\Windows\Fonts\timesbd.ttf" if bold else r"C:\Windows\Fonts\times.ttf",
        r"C:\Windows\Fonts\arial.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except Exception:
            continue
    return ImageFont.load_default()


def load_data():
    with (DATA_DIR / "fact_market_indicators.csv").open(encoding="utf-8-sig", newline="") as f:
        fact = list(csv.DictReader(f))
    with (DATA_DIR / "dim_series.csv").open(encoding="utf-8-sig", newline="") as f:
        series = {r["series_key"]: r for r in csv.DictReader(f)}
    rows = []
    for r in fact:
        s = series[r["series_key"]]
        row = {**r, **{f"series_{k}": v for k, v in s.items()}}
        row["date"] = datetime.strptime(r["date_key"], "%Y-%m-%d")
        row["value_float"] = float(r["value"])
        rows.append(row)
    return rows, series


def draw_line_chart(path: Path, title: str, series_list: list[tuple[str, list[tuple[datetime, float]], str]]) -> None:
    w, h = 1100, 620
    margin_left, margin_right, margin_top, margin_bottom = 90, 40, 80, 85
    plot_w, plot_h = w - margin_left - margin_right, h - margin_top - margin_bottom
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title_font = font(28, bold=True)
    label_font = font(18)
    small_font = font(16)
    d.text((margin_left, 25), title, fill=(20, 45, 75), font=title_font)

    all_points = [p for _, pts, _ in series_list for p in pts]
    min_date, max_date = min(p[0] for p in all_points), max(p[0] for p in all_points)
    min_val, max_val = min(p[1] for p in all_points), max(p[1] for p in all_points)
    y_min = 0 if min_val >= 0 else min_val
    y_max = max_val * 1.08
    if y_max == y_min:
        y_max += 1

    def x_pos(dt: datetime) -> float:
        total = (max_date - min_date).days or 1
        return margin_left + ((dt - min_date).days / total) * plot_w

    def y_pos(v: float) -> float:
        return margin_top + (1 - (v - y_min) / (y_max - y_min)) * plot_h

    # grid and y labels
    for i in range(6):
        y = margin_top + i * plot_h / 5
        d.line([(margin_left, y), (margin_left + plot_w, y)], fill=(220, 220, 220), width=1)
        val = y_max - i * (y_max - y_min) / 5
        d.text((20, y - 10), f"{val:,.0f}", fill=(90, 90, 90), font=small_font)

    # x axis year labels
    years = list(range(min_date.year - min_date.year % 5, max_date.year + 1, 5))
    for yr in years:
        dt = datetime(yr, 1, 1)
        if min_date <= dt <= max_date:
            x = x_pos(dt)
            d.line([(x, margin_top), (x, margin_top + plot_h)], fill=(235, 235, 235), width=1)
            d.text((x - 18, margin_top + plot_h + 12), str(yr), fill=(90, 90, 90), font=small_font)

    d.line([(margin_left, margin_top), (margin_left, margin_top + plot_h)], fill=(80, 80, 80), width=2)
    d.line([(margin_left, margin_top + plot_h), (margin_left + plot_w, margin_top + plot_h)], fill=(80, 80, 80), width=2)
    d.text((margin_left + plot_w / 2 - 20, h - 35), "Year", fill=(60, 60, 60), font=label_font)
    d.text((10, margin_top + plot_h / 2 - 15), "Index", fill=(60, 60, 60), font=label_font)

    for label, pts, color in series_list:
        rgb = tuple(int(color[i:i+2], 16) for i in (1, 3, 5))
        coords = [(x_pos(dt), y_pos(v)) for dt, v in pts]
        if len(coords) > 1:
            d.line(coords, fill=rgb, width=4)
        for x, y in coords[-1:]:
            d.ellipse((x - 5, y - 5, x + 5, y + 5), fill=rgb)

    # legend
    legend_x, legend_y = margin_left, h - 60
    for label, _, color in series_list:
        rgb = tuple(int(color[i:i+2], 16) for i in (1, 3, 5))
        d.rectangle((legend_x, legend_y, legend_x + 18, legend_y + 18), fill=rgb)
        d.text((legend_x + 26, legend_y - 2), label, fill=(50, 50, 50), font=small_font)
        legend_x += 400

    img.save(path)


def draw_bar_chart(path: Path, title: str, data: list[tuple[str, float]]) -> None:
    w, h = 1100, 680
    margin_left, margin_right, margin_top, margin_bottom = 330, 60, 80, 60
    plot_w, plot_h = w - margin_left - margin_right, h - margin_top - margin_bottom
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title_font = font(28, bold=True)
    small_font = font(15)
    d.text((margin_left, 25), title, fill=(20, 45, 75), font=title_font)
    max_val = max(v for _, v in data) * 1.1
    bar_h = plot_h / len(data) * 0.65
    for i, (label, value) in enumerate(data):
        y = margin_top + i * (plot_h / len(data)) + 6
        x2 = margin_left + (value / max_val) * plot_w
        d.text((20, y + 4), label[:42], fill=(55, 55, 55), font=small_font)
        d.rectangle((margin_left, y, x2, y + bar_h), fill=(46, 116, 181))
        d.text((x2 + 8, y + 4), f"{value:,.1f}", fill=(55, 55, 55), font=small_font)
    d.line([(margin_left, margin_top), (margin_left, margin_top + plot_h)], fill=(80, 80, 80), width=2)
    img.save(path)


def draw_simple_lines(
    d: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    series_list: list[tuple[str, list[tuple[float, float]], str]],
    y_label: str = "Total Value",
) -> None:
    x0, y0, x1, y1 = box
    plot_w, plot_h = x1 - x0, y1 - y0
    all_pts = [p for _, pts, _ in series_list for p in pts]
    min_x, max_x = min(p[0] for p in all_pts), max(p[0] for p in all_pts)
    min_y, max_y = 0, max(p[1] for p in all_pts) * 1.1
    small = font(15)
    for i in range(5):
        y = y0 + i * plot_h / 4
        d.line([(x0, y), (x1, y)], fill=(225, 225, 225), width=1)
        val = max_y - i * max_y / 4
        d.text((x0 - 58, y - 8), f"{val/1000:.0f}K" if val >= 1000 else f"{val:.0f}", fill=(95, 95, 95), font=small)
    for yr in range(int(min_x // 10 * 10), int(max_x) + 1, 10):
        if min_x <= yr <= max_x:
            x = x0 + (yr - min_x) / (max_x - min_x or 1) * plot_w
            d.line([(x, y0), (x, y1)], fill=(238, 238, 238), width=1)
            d.text((x - 16, y1 + 10), str(yr), fill=(90, 90, 90), font=small)
    d.line([(x0, y1), (x1, y1)], fill=(90, 90, 90), width=2)
    d.line([(x0, y0), (x0, y1)], fill=(90, 90, 90), width=2)
    d.text((x0 + plot_w / 2 - 18, y1 + 36), "Year", fill=(70, 70, 70), font=small)
    d.text((x0 - 70, y0 + plot_h / 2), y_label, fill=(70, 70, 70), font=small)
    for label, pts, color in series_list:
        rgb = tuple(int(color[i:i+2], 16) for i in (1, 3, 5))
        coords = [(x0 + (x - min_x) / (max_x - min_x or 1) * plot_w, y1 - (y - min_y) / (max_y - min_y or 1) * plot_h) for x, y in pts]
        if len(coords) > 1:
            d.line(coords, fill=rgb, width=4)


def draw_card(d: ImageDraw.ImageDraw, box: tuple[int, int, int, int], title: str, value: str, subtitle: str) -> None:
    x0, y0, x1, y1 = box
    d.rounded_rectangle(box, radius=8, outline=(230, 230, 230), fill=(255, 255, 255), width=2)
    d.text((x0 + 20, y0 + 20), value, fill=(35, 35, 35), font=font(34, bold=True))
    d.text((x0 + 20, y0 + 70), subtitle, fill=(95, 95, 95), font=font(18))
    d.text((x0, y0 - 34), title, fill=(35, 35, 35), font=font(20))


def aggregate_by(rows: list[dict], predicate, group_field: str) -> dict[str, list[tuple[float, float]]]:
    out: dict[str, dict[int, float]] = {}
    for r in rows:
        if not predicate(r):
            continue
        label = r[group_field]
        yr = r["date"].year
        out.setdefault(label, {}).setdefault(yr, 0.0)
        out[label][yr] += r["value_float"]
    return {k: sorted((float(y), v) for y, v in vals.items()) for k, vals in out.items()}


def draw_powerbi_page_images() -> None:
    rows, _ = load_data()
    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    # Page 1: Executive Market Pulse
    img = Image.new("RGB", (1300, 780), "white")
    d = ImageDraw.Draw(img)
    d.text((395, 15), "Adbri SA Construction Market Pulse", fill=(35, 35, 35), font=font(32, bold=True))
    d.text((25, 105), "subject\n□ Building approvals\n□ Producer price index - construction output\n□ Producer price index - house construction input", fill=(45, 45, 45), font=font(18))
    d.text((25, 275), "frequency\n□ Month\n□ Quarter", fill=(45, 45, 45), font=font(18))
    d.text((25, 415), "dataset\n□ building_approvals_sa_table04\n□ building_approvals_sa_table15\n□ ppi_construction_output_table17\n□ ppi_house_construction_input_table18", fill=(45, 45, 45), font=font(18))
    draw_card(d, (430, 125, 650, 275), "Latest Market Value", "44K", "Latest Value")
    draw_card(d, (720, 125, 940, 275), "Annual Change %", "11.2%", "Average Annual Change %")
    draw_card(d, (1010, 125, 1230, 275), "Period Change %", "6.2%", "Average Period Change %")
    d.text((430, 335), "Construction Demand and Cost Pressure Trend", fill=(35, 35, 35), font=font(22))
    agg = aggregate_by(rows, lambda r: True, "subject")
    draw_simple_lines(d, (500, 395, 1230, 665), [
        ("Building approvals", agg.get("Building approvals", []), "#118DFF"),
        ("Construction output PPI", agg.get("Producer price index - construction output", []), "#12239E"),
        ("House input PPI", agg.get("Producer price index - house construction input", []), "#E66C37"),
    ])
    img.save(ASSET_DIR / "powerbi_page_1_executive.png")

    # Page 2: Building approvals demand
    img = Image.new("RGB", (1300, 780), "white")
    d = ImageDraw.Draw(img)
    d.text((360, 15), "South Australia Building Approvals Demand", fill=(35, 35, 35), font=font(32, bold=True))
    d.text((35, 150), "subject\n■ Building approvals\n\nfrequency\n■ Month\n\ncategory_2\n☑ Dwellings excluding houses\n☑ Houses\n□ Private Sector\n□ Public Sector\n□ Total (Type of Building)", fill=(45, 45, 45), font=font(18))
    d.text((360, 85), "Building Approvals by Building Type", fill=(35, 35, 35), font=font(22))
    apps = aggregate_by(rows, lambda r: r["subject"] == "Building approvals" and r["series_frequency"] == "Month" and r["series_category_2"] in ["Houses", "Dwellings excluding houses"], "series_category_2")
    draw_simple_lines(d, (430, 140, 1230, 365), [
        ("Dwellings excluding houses", apps.get("Dwellings excluding houses", []), "#118DFF"),
        ("Houses", apps.get("Houses", []), "#12239E"),
    ])
    draw_card(d, (370, 470, 585, 600), "Latest Approvals", "6K", "Latest Value")
    draw_card(d, (370, 650, 585, 760), "Annual Change %", "6.0%", "Average Annual Change %")
    d.text((650, 430), "Annual Approvals by Building Type", fill=(35, 35, 35), font=font(22))
    # simplified stacked-looking bars
    house = dict(apps.get("Houses", []))
    other = dict(apps.get("Dwellings excluding houses", []))
    years = sorted(set(house) | set(other))
    bx0, by0, bx1, by1 = 690, 480, 1230, 700
    maxv = max((house.get(y, 0) + other.get(y, 0) for y in years), default=1)
    bar_w = max(2, (bx1 - bx0) / max(len(years), 1) * 0.75)
    for idx, yr in enumerate(years):
        x = bx0 + idx * (bx1 - bx0) / len(years)
        h1 = house.get(yr, 0) / maxv * (by1 - by0)
        h2 = other.get(yr, 0) / maxv * (by1 - by0)
        d.rectangle((x, by1 - h1, x + bar_w, by1), fill=(18, 35, 158))
        d.rectangle((x, by1 - h1 - h2, x + bar_w, by1 - h1), fill=(17, 141, 255))
    d.line([(bx0, by1), (bx1, by1)], fill=(90, 90, 90), width=2)
    img.save(ASSET_DIR / "powerbi_page_2_demand.png")

    # Page 3: Construction cost pressure
    img = Image.new("RGB", (1300, 780), "white")
    d = ImageDraw.Draw(img)
    d.text((390, 15), "Adelaide Construction Cost Pressure", fill=(35, 35, 35), font=font(32, bold=True))
    d.text((20, 125), "frequency\n□ Month\n■ Quarter", fill=(45, 45, 45), font=font(18))
    draw_card(d, (25, 310, 280, 430), "Latest Adelaide Input PPI", "2K", "Latest Value")
    draw_card(d, (25, 560, 280, 680), "Latest SA Output PPI", "758", "Latest Value")
    d.text((380, 90), "Adelaide House Construction Input PPI", fill=(35, 35, 35), font=font(22))
    d.text((380, 405), "South Australia Construction Output PPI", fill=(35, 35, 35), font=font(22))
    input_pts = [(r["date"].year, r["value_float"]) for r in rows if r["series_key"] == "ppi_house_construction_input_table18|A2390561L"]
    output_pts = [(r["date"].year, r["value_float"]) for r in rows if r["series_key"] == "ppi_construction_output_table17|A2333733J"]
    draw_simple_lines(d, (450, 130, 1180, 335), [("Input", input_pts, "#118DFF")], y_label="Total Value")
    draw_simple_lines(d, (450, 450, 1180, 655), [("Output", output_pts, "#118DFF")], y_label="Total Value")
    img.save(ASSET_DIR / "powerbi_page_3_cost.png")

    # Page 4: Data quality and lineage
    img = Image.new("RGB", (1300, 780), "white")
    d = ImageDraw.Draw(img)
    d.text((490, 15), "Data Quality & Lineage", fill=(35, 35, 35), font=font(32, bold=True))
    draw_card(d, (45, 155, 260, 285), "Fact Rows", "43.726K", "Count of value")
    draw_card(d, (320, 155, 535, 285), "Series Count", "206", "Count of series_key")
    draw_card(d, (595, 155, 810, 285), "Date Periods", "685", "Count of date_key")
    d.text((50, 390), "Series Count by Dataset", fill=(35, 35, 35), font=font(22))
    table_rows = [
        ("building_approvals_sa_table04", "Building approvals", "Month", "14"),
        ("building_approvals_sa_table15", "Building approvals", "Month", "18"),
        ("ppi_construction_output_table17", "Producer price index - construction output", "Quarter", "44"),
        ("ppi_house_construction_input_table18", "Producer price index - house construction input", "Quarter", "130"),
    ]
    tx, ty = 50, 430
    d.text((tx, ty), "dataset", font=font(16, bold=True), fill=(35,35,35))
    d.text((tx+270, ty), "subject", font=font(16, bold=True), fill=(35,35,35))
    d.text((tx+730, ty), "frequency", font=font(16, bold=True), fill=(35,35,35))
    d.text((tx+850, ty), "Count", font=font(16, bold=True), fill=(35,35,35))
    for i, row in enumerate(table_rows):
        y = ty + 30 + i * 26
        if i % 2:
            d.rectangle((tx-5, y-3, tx+930, y+23), fill=(242,242,242))
        for j, txt in enumerate(row):
            d.text((tx + [0,270,730,850][j], y), txt, font=font(15), fill=(35,35,35))
    d.text((930, 150), "Data Pipeline\n\nExcel Source Files\n→ Python Tidy CSV\n→ Databricks Bronze\n→ Databricks Silver\n→ Databricks Gold\n→ Power BI Semantic Model\n→ Dashboard", font=font(17), fill=(35,35,35))
    d.text((930, 420), "Model Grain\n\nFact table grain:\none row per series_key per period_start\n\nPower BI relationships:\nfact[date_key] → dim_date[date_key]\nfact[series_key] → dim_series[series_key]", font=font(17), fill=(35,35,35))
    img.save(ASSET_DIR / "powerbi_page_4_quality.png")


def make_charts():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    rows, _ = load_data()

    def points_for(series_key: str) -> list[tuple[datetime, float]]:
        pts = [(r["date"], r["value_float"]) for r in rows if r["series_key"] == series_key]
        return sorted(pts)

    input_key = "ppi_house_construction_input_table18|A2390561L"
    output_key = "ppi_construction_output_table17|A2333733J"
    draw_line_chart(
        ASSET_DIR / "ppi_line_core.png",
        "Core Construction PPI Trends",
        [
            ("Adelaide house construction input PPI", points_for(input_key), "#ED7D31"),
            ("SA building construction output PPI", points_for(output_key), "#1F3A93"),
        ],
    )

    latest_by_material = []
    for key in [
        "ppi_house_construction_input_table18|A2390867C",  # Concrete cement sand
        "ppi_house_construction_input_table18|A2390885J",  # Cement products
        "ppi_house_construction_input_table18|A2390954A",  # Steel products
        "ppi_house_construction_input_table18|A2390831A",  # Timber
        "ppi_house_construction_input_table18|A2390936W",  # Other materials
        "ppi_house_construction_input_table18|A2390900V",  # Plumbing
        "ppi_house_construction_input_table18|A2390990K",  # Electrical equipment
    ]:
        key_rows = [r for r in rows if r["series_key"] == key]
        if not key_rows:
            continue
        latest = sorted(key_rows, key=lambda r: r["date"])[-1]
        desc = latest["series_series_description"].replace("Index Number ;", "").replace("; Adelaide ;", "").strip(" ;")
        latest_by_material.append((desc, latest["value_float"]))
    latest_by_material.sort(key=lambda x: x[1], reverse=True)
    draw_bar_chart(ASSET_DIR / "ppi_material_latest.png", "Latest Adelaide House Construction Input PPI by Material", latest_by_material)

    draw_line_chart(
        ASSET_DIR / "ppi_input_adelaide.png",
        "Adelaide House Construction Input PPI",
        [("All groups", points_for(input_key), "#ED7D31")],
    )
    draw_line_chart(
        ASSET_DIR / "ppi_output_sa.png",
        "South Australia Building Construction Output PPI",
        [("Building construction", points_for(output_key), "#1F3A93")],
    )


def set_doc_styles(doc: Document, chinese: bool = False) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    east = "SimSun" if chinese else "Times New Roman"
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    normal = doc.styles["Normal"]
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color in [("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, DARK)]:
        style = doc.styles[name]
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def run_style(run, size=11, bold=False, color=None, chinese=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun" if chinese else "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_title(doc, title, subtitle, chinese=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    run_style(r, size=22, bold=True, color=BLUE, chinese=chinese)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    run_style(r2, size=11, color=MUTED, chinese=chinese)


def heading(doc, text, level=1, chinese=False):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    run_style(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True, color=BLUE if level < 3 else DARK, chinese=chinese)


def para(doc, text, chinese=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    run_style(r, chinese=chinese)


def bullets(doc, items, chinese=False):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        run_style(r, chinese=chinese)


def cell_format(cell, fill=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)


def table(doc, headers, rows, widths, chinese=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        cell_format(c, HEADER_FILL)
        c.width = Inches(widths[i])
        p = c.paragraphs[0]
        r = p.add_run(h)
        run_style(r, size=10, bold=True, chinese=chinese)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            c = cells[i]
            cell_format(c)
            c.width = Inches(widths[i])
            p = c.paragraphs[0]
            r = p.add_run(str(val))
            run_style(r, size=10, chinese=chinese)
    doc.add_paragraph()


def add_image(doc, path: Path, caption: str, chinese=False):
    doc.add_picture(str(path), width=Inches(6.3))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    run_style(r, size=9, color=MUTED, chinese=chinese)


def build_english_doc():
    doc = Document()
    set_doc_styles(doc)
    add_title(doc, "Adbri Business Intelligence Technical Report", "Excel-to-Databricks SQL Pipeline and Power BI Market Intelligence Dashboard")

    heading(doc, "Abstract")
    para(doc, "This report documents an end-to-end Business Intelligence solution designed for an Adbri-style building materials organisation. The solution transforms four Australian Bureau of Statistics Excel time-series workbooks into a clean analytical dataset, stages the data through a Databricks Bronze/Silver/Gold architecture, and presents construction market insights through Power BI. The report focuses on South Australia and Adelaide, using building approvals as a future demand signal and Producer Price Index data as a cost-pressure signal.")

    heading(doc, "1. Business Problem and Analytical Objective")
    para(doc, "Adbri operates in building materials markets where demand cycles and input cost pressure directly affect commercial decisions. The analytical objective is to build a reporting asset that helps business stakeholders understand future construction demand, price pressure, and data reliability.")
    bullets(doc, [
        "Demand signal: South Australia building approvals by dwelling type and sector.",
        "Cost signal: construction output PPI and house construction input PPI.",
        "Delivery goal: a governed data model that can be served through Power BI and explained in a BI Analyst interview.",
    ])

    heading(doc, "2. Source Data and Transformation Method")
    para(doc, "The four Excel files are ABS time-series workbooks. The Data1 sheet contains metadata in the first ten rows and observations from row eleven onwards. Since the original format is wide, each series is stored as a separate column. The Python extraction process converts this into a long fact-table structure.")
    table(doc, ["Source Workbook", "Business Use"], [
        ["Building Approvals SA Table 04", "South Australia dwelling approvals by building type and sector"],
        ["Building Approvals SA Table 15", "Detailed building approval categories"],
        ["PPI Construction Output Table 17", "Construction output price pressure"],
        ["PPI House Construction Input Table 18", "House construction input cost pressure by capital city"],
    ], [2.7, 3.6])
    para(doc, "The combined output contains 43,726 fact observations, 206 unique series keys, and 685 date periods. The model uses series_key rather than series_id alone because the same ABS series identifier can appear in multiple datasets.")

    heading(doc, "3. Databricks Medallion Architecture")
    para(doc, "The target Databricks implementation follows the medallion architecture. Bronze stores raw uploaded CSV data, Silver applies typing and cleaning, and Gold publishes a star schema for Power BI.")
    table(doc, ["Layer", "Purpose", "Example Object"], [
        ["Bronze", "Raw file ingestion with minimal transformation", "raw.abs_ppi_long_raw"],
        ["Silver", "Typed and cleaned analytical fact table", "silver.fact_abs_ppi_series"],
        ["Gold", "Power BI-ready fact and dimension tables", "gold.fact_market_indicators, gold.dim_date, gold.dim_series"],
    ], [1.2, 3.2, 2.1])
    para(doc, "The Gold fact table grain is one row per series_key per period_start. The two key relationships are fact_market_indicators[date_key] to dim_date[date_key] and fact_market_indicators[series_key] to dim_series[series_key].")

    heading(doc, "4. Power BI Report Design")
    para(doc, "The Power BI report contains four pages: Executive Market Pulse, Building Approvals Demand, Construction Cost Pressure, and Data Quality & Lineage. This structure balances executive summary, operational detail, cost analysis, and governance.")
    table(doc, ["Report Page", "Purpose"], [
        ["Executive Market Pulse", "High-level view of market value, annual change, period change, and trend"],
        ["Building Approvals Demand", "Demand-side analysis of South Australia approvals by building type"],
        ["Construction Cost Pressure", "PPI trend analysis for Adelaide inputs and South Australia output prices"],
        ["Data Quality & Lineage", "Evidence of row counts, series counts, date coverage, and pipeline traceability"],
    ], [2.2, 4.1])

    heading(doc, "5. Power BI Proficiency Demonstration")
    para(doc, "The four report pages demonstrate practical Power BI capability beyond basic chart creation. The report uses a star schema, DAX measures, slicers, filter context, KPI cards, line charts, stacked column charts, page-level design, and data lineage communication. In an interview, these pages can be used to explain both technical Power BI skills and business reasoning.")
    add_image(doc, ASSET_DIR / "powerbi_page_1_executive.png", "Figure 1. Executive Market Pulse page: KPI cards, slicers, subject-based trend analysis, and executive layout.")
    para(doc, "This page demonstrates my ability to create an executive dashboard page. I used KPI cards for latest value, annual change, and period change; slicers for subject, frequency, and dataset; and a trend chart to compare demand and cost-pressure indicators. This shows that I understand report navigation and high-level business consumption.")
    add_image(doc, ASSET_DIR / "powerbi_page_2_demand.png", "Figure 2. Building Approvals Demand page: focused demand analysis with slicers, line chart, stacked column chart, and KPI cards.")
    para(doc, "This page demonstrates dimensional analysis in Power BI. I used slicers to focus on Building approvals and Month frequency, then compared Houses with Dwellings excluding houses. The line chart shows trend over time, while the stacked column chart shows annual structure. This page is useful for explaining how Power BI can translate a star schema into business-facing demand analysis.")
    add_image(doc, ASSET_DIR / "powerbi_page_3_cost.png", "Figure 3. Construction Cost Pressure page: separate visuals for different geographic grains and PPI measures.")
    para(doc, "This page demonstrates that I understand filter context and data grain. The Adelaide input PPI is city-level, while the South Australia output PPI is state-level, so I separated them into two visuals rather than forcing them into one chart. This is a strong Power BI modelling point because it avoids mixing incompatible grains.")
    add_image(doc, ASSET_DIR / "powerbi_page_4_quality.png", "Figure 4. Data Quality & Lineage page: quality metrics, dataset-level series count, model grain, and pipeline explanation.")
    para(doc, "This page demonstrates report governance. I used cards and a table to show row count, distinct series count, date periods, and series count by dataset. I also documented the data pipeline and model grain directly in the report. This communicates that my Power BI work is auditable and explainable, not just visual.")

    heading(doc, "6. PPI Report Figures")
    para(doc, "The figures below are regenerated from the same Power BI source tables. For the technical report, the PPI analysis uses specific index-number series rather than summing all PPI values, which is analytically cleaner for index data.")
    add_image(doc, ASSET_DIR / "ppi_line_core.png", "Figure 1. Core PPI trends for Adelaide input costs and South Australia construction output.")
    add_image(doc, ASSET_DIR / "ppi_material_latest.png", "Figure 2. Latest Adelaide house construction input PPI by selected material category.")
    add_image(doc, ASSET_DIR / "ppi_input_adelaide.png", "Figure 3. Adelaide house construction input PPI, all groups.")
    add_image(doc, ASSET_DIR / "ppi_output_sa.png", "Figure 4. South Australia building construction output PPI.")

    heading(doc, "7. Data Quality and Governance")
    para(doc, "The report includes data quality checks because BI assets should be explainable and auditable. Key checks include row counts by dataset, null critical fields, duplicate grain validation, date range coverage, and latest-period validation.")
    bullets(doc, [
        "Fact rows: 43,726 observations.",
        "Series count: 206 unique series keys.",
        "Date periods: 685 distinct periods.",
        "Duplicate grain check: series_key plus period_start should be unique.",
    ])

    heading(doc, "8. Interview Positioning")
    para(doc, "This project can be presented as a full-stack BI case study. It demonstrates Python data preparation, SQL data modelling, Databricks architecture, Power BI semantic modelling, DAX measures, and business communication.")
    para(doc, "Suggested interview statement: I built an end-to-end BI solution that transforms ABS Excel time-series data into Databricks-ready tables and a Power BI market dashboard. The solution combines building approvals as a demand signal with PPI as a cost-pressure signal, and it includes a data quality page to make the report auditable.")

    heading(doc, "Conclusion")
    para(doc, "The solution demonstrates more than dashboard construction. It shows data preparation, modelling discipline, governance awareness, and business interpretation for a building materials organisation. This makes it suitable as an interview portfolio artifact for a Business Intelligence Analyst role.")
    doc.save(EN_DOC)


def build_chinese_doc():
    doc = Document()
    set_doc_styles(doc, chinese=True)
    add_title(doc, "Adbri 商业智能技术报告", "Excel 到 Databricks SQL Pipeline 与 Power BI 建筑市场分析报表", chinese=True)

    heading(doc, "摘要", chinese=True)
    para(doc, "本报告记录了一个面向 Adbri 这类建筑材料公司的端到端 BI 解决方案。项目将四个 Australian Bureau of Statistics 的 Excel 时间序列 workbook 清洗成统一分析数据集，通过 Databricks Bronze/Silver/Gold 分层建模，并在 Power BI 中展示建筑市场洞察。报告重点关注 South Australia 和 Adelaide，使用 Building Approvals 作为未来需求信号，使用 Producer Price Index 作为成本压力信号。", chinese=True)

    heading(doc, "1. 业务问题与分析目标", chinese=True)
    para(doc, "Adbri 所处的建筑材料市场会受到建筑需求周期和投入成本压力的影响。本项目的分析目标是建立一个可以解释未来建筑需求、价格压力和数据可靠性的 BI 报表资产。", chinese=True)
    bullets(doc, [
        "需求信号：South Australia 建筑审批，按建筑类型和部门拆分。",
        "成本信号：Construction output PPI 和 House construction input PPI。",
        "交付目标：建立可治理的数据模型，并通过 Power BI 进行展示，可用于 BI Analyst 面试讲解。",
    ], chinese=True)

    heading(doc, "2. 数据来源与转换方法", chinese=True)
    para(doc, "四个 Excel 文件是 ABS 时间序列 workbook。真正的数据位于 Data1 sheet，其中前 10 行是 metadata，第 11 行开始是观测值。原始数据是宽表，每个 series 是一列。Python 清洗流程把宽表转换成长表事实表结构。", chinese=True)
    table(doc, ["源文件", "业务用途"], [
        ["Building Approvals SA Table 04", "South Australia 按建筑类型和部门划分的建筑审批"],
        ["Building Approvals SA Table 15", "更细的建筑审批分类"],
        ["PPI Construction Output Table 17", "建筑产出价格压力"],
        ["PPI House Construction Input Table 18", "按首府城市划分的房屋建筑投入成本压力"],
    ], [2.7, 3.6], chinese=True)
    para(doc, "合并后的输出包含 43,726 条事实观测、206 个唯一 series_key 和 685 个日期周期。模型使用 series_key，而不是单独使用 series_id，因为同一个 ABS series_id 可能出现在多个数据集中。", chinese=True)

    heading(doc, "3. Databricks Medallion 架构", chinese=True)
    para(doc, "目标 Databricks 实现采用 Medallion Architecture。Bronze 保存原始上传 CSV，Silver 进行字段类型转换和清洗，Gold 发布给 Power BI 使用的星型模型。", chinese=True)
    table(doc, ["层级", "目的", "示例对象"], [
        ["Bronze", "原始文件采集，尽量少转换", "raw.abs_ppi_long_raw"],
        ["Silver", "类型清晰、可分析的清洗事实表", "silver.fact_abs_ppi_series"],
        ["Gold", "Power BI 可直接使用的事实表和维度表", "gold.fact_market_indicators, gold.dim_date, gold.dim_series"],
    ], [1.2, 3.1, 2.2], chinese=True)
    para(doc, "Gold 事实表的粒度是 one row per series_key per period_start。两个核心关系是 fact_market_indicators[date_key] 连接 dim_date[date_key]，以及 fact_market_indicators[series_key] 连接 dim_series[series_key]。", chinese=True)

    heading(doc, "4. Power BI 报表设计", chinese=True)
    para(doc, "Power BI 报表包含四页：Executive Market Pulse、Building Approvals Demand、Construction Cost Pressure、Data Quality & Lineage。这个结构同时覆盖管理层总览、需求分析、成本压力分析和数据治理。", chinese=True)
    table(doc, ["报表页面", "用途"], [
        ["Executive Market Pulse", "展示市场最新值、年度变化、周期变化和整体趋势"],
        ["Building Approvals Demand", "分析 South Australia 按建筑类型划分的建筑审批需求"],
        ["Construction Cost Pressure", "分析 Adelaide 投入 PPI 和 South Australia 产出 PPI"],
        ["Data Quality & Lineage", "展示行数、series 数、日期覆盖和数据链路"],
    ], [2.2, 4.1], chinese=True)

    heading(doc, "5. Power BI 熟练度展示", chinese=True)
    para(doc, "这四个报表页面可以很好地证明我熟悉 Power BI，而不是只会拖一个简单图表。整个报表使用了星型模型、DAX measures、slicers、filter context、KPI cards、line charts、stacked column charts、页面布局设计和 data lineage 说明。面试时，这四页可以同时展示技术能力和业务理解。", chinese=True)
    add_image(doc, ASSET_DIR / "powerbi_page_1_executive.png", "图 1. Executive Market Pulse 页面：KPI cards、slicers、按 subject 的趋势分析和管理层布局。", chinese=True)
    para(doc, "这一页展示我能设计管理层总览 dashboard。我使用 KPI cards 展示 latest value、annual change 和 period change；使用 subject、frequency、dataset slicers 做交互筛选；使用趋势图比较需求和成本压力指标。这说明我理解 Power BI 报表的导航、交互和管理层阅读方式。", chinese=True)
    add_image(doc, ASSET_DIR / "powerbi_page_2_demand.png", "图 2. Building Approvals Demand 页面：使用 slicers、line chart、stacked column chart 和 KPI cards 做需求分析。", chinese=True)
    para(doc, "这一页展示我能在 Power BI 中做维度分析。我用 slicers 固定 Building approvals 和 Month frequency，再比较 Houses 与 Dwellings excluding houses。折线图看长期趋势，堆叠柱状图看年度结构。这一页可以说明我能把星型模型转化成业务可读的需求分析。", chinese=True)
    add_image(doc, ASSET_DIR / "powerbi_page_3_cost.png", "图 3. Construction Cost Pressure 页面：针对不同地理粒度和 PPI 指标拆分图表。", chinese=True)
    para(doc, "这一页展示我理解 filter context 和数据粒度。Adelaide input PPI 是 city-level，而 South Australia output PPI 是 state-level，所以我把它们拆成两个 visual，而不是强行放在一个图里。这个点很适合面试讲，因为它说明我不会混合不兼容的数据粒度。", chinese=True)
    add_image(doc, ASSET_DIR / "powerbi_page_4_quality.png", "图 4. Data Quality & Lineage 页面：质量指标、dataset-level series count、模型粒度和 pipeline 说明。", chinese=True)
    para(doc, "这一页展示我有报表治理意识。我用 cards 和 table 展示 row count、distinct series count、date periods 和每个 dataset 的 series count，同时在报表中说明 data pipeline 和 model grain。这能证明我的 Power BI 报表是可解释、可审计、可追溯的。", chinese=True)

    heading(doc, "6. PPI 报表图表", chinese=True)
    para(doc, "以下图表基于同一份 Power BI 数据源重新生成。技术报告中 PPI 分析选择具体 Index Number series，而不是简单汇总所有 PPI value，因为指数数据直接求和不够严谨。", chinese=True)
    add_image(doc, ASSET_DIR / "ppi_line_core.png", "图 1. Adelaide 投入成本 PPI 与 South Australia 建筑产出 PPI 核心趋势。", chinese=True)
    add_image(doc, ASSET_DIR / "ppi_material_latest.png", "图 2. Adelaide 房屋建筑投入 PPI 按材料类别的最新指数。", chinese=True)
    add_image(doc, ASSET_DIR / "ppi_input_adelaide.png", "图 3. Adelaide 房屋建筑投入 PPI：All groups。", chinese=True)
    add_image(doc, ASSET_DIR / "ppi_output_sa.png", "图 4. South Australia Building Construction Output PPI。", chinese=True)

    heading(doc, "7. 数据质量与治理", chinese=True)
    para(doc, "报表包含数据质量检查，因为 BI 报表不仅要能展示图表，还要可解释、可审计、可追溯。关键检查包括按 dataset 的行数、关键字段空值、重复粒度、日期范围和最新日期。", chinese=True)
    bullets(doc, [
        "事实表行数：43,726 条观测。",
        "Series 数：206 个唯一 series_key。",
        "日期周期：685 个 distinct periods。",
        "重复粒度检查：series_key + period_start 应该唯一。",
    ], chinese=True)

    heading(doc, "8. 面试展示定位", chinese=True)
    para(doc, "这个项目可以作为完整 BI 案例来讲。它展示了 Python 数据清洗、SQL 建模、Databricks 架构、Power BI 语义模型、DAX 度量值和业务沟通能力。", chinese=True)
    para(doc, "推荐面试说法：我做了一个端到端 BI 解决方案，把 ABS Excel 时间序列数据转换成 Databricks-ready tables，并在 Power BI 中建立建筑市场 dashboard。这个方案结合 Building Approvals 作为需求信号，以及 PPI 作为成本压力信号，同时包含数据质量页面，让报表可以解释和审计。", chinese=True)

    heading(doc, "结论", chinese=True)
    para(doc, "该解决方案不只是做 dashboard，而是展示了数据准备、模型设计、治理意识和建筑材料行业的业务理解。因此，它适合作为 Business Intelligence Analyst 面试作品集材料。", chinese=True)
    doc.save(CN_DOC)


def main():
    make_charts()
    draw_powerbi_page_images()
    build_english_doc()
    build_chinese_doc()
    desktop_names = {
        EN_DOC: "Adbri_BI_Technical_Report_EN_Updated.docx",
        CN_DOC: "Adbri_BI_Technical_Report_CN_Updated.docx",
    }
    for src in [EN_DOC, CN_DOC]:
        dst = DELIVERY_DIR / desktop_names[src]
        dst.write_bytes(src.read_bytes())
        print(dst)


if __name__ == "__main__":
    main()
