from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
OUTPUT_PATH = BASE_DIR / "docs" / "Adbri_BI_Report_Walkthrough_CN_EN.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY_FILL = "F2F4F7"


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.color.rgb = BLUE if level <= 2 else DARK_BLUE


def add_para(doc: Document, text: str, bold_label: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_label and text.startswith(bold_label):
        r1 = p.add_run(bold_label)
        r1.bold = True
        r2 = p.add_run(text[len(bold_label):])
        runs = [r1, r2]
    else:
        runs = [p.add_run(text)]
    for run in runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_fill(cell, GRAY_FILL)
        set_cell_text(cell, header, bold=True)
        cell.width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            cells[idx].width = Inches(widths[idx])
    doc.add_paragraph()


def build_doc() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Adbri BI Report Walkthrough")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("中英文面试讲解稿 | SQL, Databricks, Power BI, Python, Fabric, Azure")
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(85, 85, 85)

    add_heading(doc, "1. One-Minute Overview | 一分钟总览", 1)
    add_para(
        doc,
        "English: This report is a construction market intelligence dashboard for an Adbri-style building materials business. It combines South Australia building approvals as a demand signal and Producer Price Index data as a cost pressure signal. I used Python to reshape raw ABS Excel files, prepared Databricks-ready tables, modelled the data as a Power BI star schema, and built four report pages for executive insight, demand analysis, cost pressure, and data quality.",
        "English:",
    )
    add_para(
        doc,
        "中文：这个报表是为 Adbri 这类建筑材料公司做的建筑市场 BI Dashboard。它结合了 South Australia 建筑审批数据作为未来需求信号，以及 Producer Price Index 作为成本压力信号。我用 Python 把 ABS 原始 Excel 清洗成长表，再准备成 Databricks 可以读取的数据结构，最后在 Power BI 里建立星型模型，做出总览、需求分析、成本压力、数据质量四个页面。",
        "中文：",
    )

    add_heading(doc, "2. Business Context | 业务背景", 1)
    add_para(
        doc,
        "English: For a building materials company, demand and cost pressure are both important. Building approvals help estimate future construction activity, while PPI indicates whether construction input and output prices are rising. Together, these indicators help commercial teams understand market direction in South Australia and Adelaide.",
        "English:",
    )
    add_para(
        doc,
        "中文：对建筑材料公司来说，需求和成本压力都很重要。Building Approvals 可以帮助判断未来建筑活动和材料需求；PPI 可以判断建筑投入成本和产出价格是否上升。两个指标结合，可以帮助业务团队理解 South Australia 和 Adelaide 市场走势。",
        "中文：",
    )

    add_table(
        doc,
        ["Dataset", "Business Meaning | 业务含义"],
        [
            ["Building Approvals SA Table 04", "Demand trend by building type and sector | 建筑类型和部门维度的需求趋势"],
            ["Building Approvals SA Table 15", "More detailed approval categories | 更细的审批分类"],
            ["PPI Construction Output Table 17", "Construction output price pressure | 建筑产出价格压力"],
            ["PPI House Construction Input Table 18", "House construction input cost pressure by city | 各城市房屋建筑投入成本压力"],
        ],
        [2.6, 3.8],
    )

    add_heading(doc, "3. Data Model | 数据模型", 1)
    add_para(doc, "English: The Power BI model uses a star schema. The central fact table is fact_market_indicators, connected to dim_date and dim_series.", "English:")
    add_para(doc, "中文：Power BI 模型采用星型模型。中间是 fact_market_indicators 事实表，两边连接 dim_date 日期维度和 dim_series 指标维度。", "中文：")
    add_bullets(
        doc,
        [
            "Fact grain: one row per series_key per period_start.",
            "Relationship 1: fact_market_indicators[date_key] -> dim_date[date_key].",
            "Relationship 2: fact_market_indicators[series_key] -> dim_series[series_key].",
            "Important modelling choice: series_key = dataset + series_id, because the same ABS series_id can appear across tables.",
        ],
    )
    add_para(
        doc,
        "中文重点：事实表粒度是一条 series_key 在一个 period_start 只有一行。这里不用单独的 series_id 做关系，而是用 dataset + series_id 生成 series_key，因为 ABS 的同一个 series_id 可能跨表重复出现。",
    )

    add_heading(doc, "4. Page-by-Page Walkthrough | 四页报表讲解", 1)

    add_heading(doc, "Page 1: Executive Market Pulse | 管理层市场总览", 2)
    add_para(
        doc,
        "English: This page gives a quick executive view of the construction market. It shows latest market value, annual change, period change, and a trend line across demand and cost pressure indicators. The slicers allow users to switch between subject, frequency, and dataset.",
        "English:",
    )
    add_para(
        doc,
        "中文：第一页是给管理层看的总览页。它展示最新值、年度变化、周期变化，以及需求和成本压力的趋势线。左侧 slicer 可以按 subject、frequency、dataset 筛选。",
        "中文：",
    )
    add_para(doc, "Interview line: This page helps executives quickly see whether the market is expanding, softening, or facing cost pressure.", "Interview line:")
    add_para(doc, "中文讲法：这一页帮助管理层快速判断市场是在增长、放缓，还是成本压力正在上升。", "中文讲法：")

    add_heading(doc, "Page 2: Building Approvals Demand | 建筑审批需求分析", 2)
    add_para(
        doc,
        "English: This page focuses on South Australia building approvals as a demand proxy. I split approvals into houses and dwellings excluding houses so the business can see which building type is driving construction demand.",
        "English:",
    )
    add_para(
        doc,
        "中文：第二页专门分析 South Australia 的建筑审批数据，把它作为未来建筑需求的代理指标。我把审批量拆成 Houses 和 Dwellings excluding houses，方便看是哪类建筑在推动需求。",
        "中文：",
    )
    add_bullets(
        doc,
        [
            "Line chart: long-term approvals trend by building type.",
            "Stacked column chart: annual approvals by building type.",
            "KPI cards: latest approvals and annual change percentage.",
        ],
    )
    add_para(doc, "Interview line: Building approvals are useful for Adbri because they are an early indicator of future demand for cement, concrete, aggregates, lime, and masonry products.", "Interview line:")
    add_para(doc, "中文讲法：对 Adbri 来说，建筑审批是未来材料需求的前置信号，因为审批上升通常意味着未来水泥、混凝土、骨料、石灰、砌块等材料需求可能上升。", "中文讲法：")

    add_heading(doc, "Page 3: Construction Cost Pressure | 建筑成本压力", 2)
    add_para(
        doc,
        "English: This page focuses on construction cost pressure using PPI data. I separated Adelaide house construction input PPI from South Australia construction output PPI because the datasets have different geographic grains: one is city-level and the other is state-level.",
        "English:",
    )
    add_para(
        doc,
        "中文：第三页分析建筑成本压力。我把 Adelaide house construction input PPI 和 South Australia construction output PPI 分成两个图，因为它们的数据地理粒度不同：一个是 city-level，一个是 state-level。",
        "中文：",
    )
    add_bullets(
        doc,
        [
            "Adelaide House Construction Input PPI: local input cost pressure.",
            "South Australia Construction Output PPI: state construction output price trend.",
            "KPI cards: latest input PPI and latest output PPI.",
        ],
    )
    add_para(doc, "Interview line: Separating these charts avoids mixing different grains and makes the analysis more reliable.", "Interview line:")
    add_para(doc, "中文讲法：把两个图分开，是为了避免混合不同粒度的数据，这样分析更严谨。", "中文讲法：")

    add_heading(doc, "Page 4: Data Quality & Lineage | 数据质量与血缘", 2)
    add_para(
        doc,
        "English: This page demonstrates data governance thinking. It shows fact row count, unique series count, date period count, series count by dataset, and the pipeline from Excel source files to Power BI.",
        "English:",
    )
    add_para(
        doc,
        "中文：第四页展示数据治理和数据质量思维。它展示事实表行数、唯一 series 数、日期周期数、每个 dataset 的 series 数，以及从 Excel 到 Power BI 的数据链路。",
        "中文：",
    )
    add_bullets(
        doc,
        [
            "Fact Rows: 43,726 observations.",
            "Series Count: 206 unique series keys.",
            "Date Periods: 685 time periods.",
            "Pipeline: Excel source files -> Python tidy CSV -> Databricks Bronze -> Silver -> Gold -> Power BI semantic model -> Dashboard.",
        ],
    )
    add_para(doc, "Interview line: I included a quality and lineage page because BI reports should be explainable and auditable, not just visually attractive.", "Interview line:")
    add_para(doc, "中文讲法：我加入数据质量和血缘页面，是因为 BI 报表不仅要好看，还要可解释、可审计、可追溯。", "中文讲法：")

    add_heading(doc, "5. Technical Architecture | 技术架构讲法", 1)
    add_table(
        doc,
        ["Skill", "How to Explain It | 如何讲"],
        [
            ["Python", "I used Python to parse ABS Excel workbooks and reshape wide time-series data into tidy long CSV. | 用 Python 清洗 ABS Excel 宽表，转成长表 CSV。"],
            ["SQL", "I used SQL concepts for Bronze/Silver/Gold modelling, joins, window calculations, and quality checks. | 用 SQL 思路做分层建模、关联、窗口计算和质量检查。"],
            ["Databricks", "The project is designed for Databricks Volume ingestion and Delta-style Bronze, Silver, Gold tables. | 设计成 Databricks Volume 上传，再建 Bronze/Silver/Gold 表。"],
            ["Power BI", "I built a star schema, measures, slicers, KPI cards, and report pages for business users. | 在 Power BI 建星型模型、度量值、筛选器、KPI 和报表页面。"],
            ["Fabric", "The same design can be implemented in Fabric Lakehouse and Power BI semantic model. | 同样架构可以迁移到 Fabric Lakehouse 和 Power BI semantic model。"],
            ["Azure", "In production, raw files can land in ADLS Gen2 and be processed by Azure Databricks with governed access. | 生产环境可以用 ADLS Gen2 存原始文件，用 Azure Databricks 处理，并做好权限治理。"],
        ],
        [1.15, 5.15],
    )

    add_heading(doc, "6. Full Interview Script | 完整面试话术", 1)
    add_heading(doc, "English Version", 2)
    add_para(
        doc,
        "I built an end-to-end BI report for an Adbri-style building materials business. The goal was to understand construction market demand and cost pressure in South Australia and Adelaide. I used ABS building approvals as a demand indicator and construction PPI as a price and cost pressure indicator. The original files were Excel workbooks with metadata rows and wide time-series columns, so I used Python to reshape them into tidy long tables.",
    )
    add_para(
        doc,
        "The target architecture is Databricks plus Power BI. In Databricks, the data can be loaded into a Bronze table, transformed into a typed Silver fact table, and published as Gold fact and dimension tables. In Power BI, I created a star schema with a market indicators fact table, a date dimension, and a series dimension. I also created DAX measures for latest value, annual change, period change, building approvals, and PPI index.",
    )
    add_para(
        doc,
        "The report has four pages: an executive market pulse page, a building approvals demand page, a construction cost pressure page, and a data quality and lineage page. The important modelling decision was using series_key instead of series_id because the same ABS series id can appear across different tables. This makes the relationships more reliable. Overall, the project shows that I can connect business context, data engineering, SQL modelling, Power BI reporting, and data quality thinking.",
    )

    add_heading(doc, "中文版本", 2)
    add_para(
        doc,
        "我做了一个端到端 BI 报表项目，场景是 Adbri 这种建筑材料公司。项目目标是分析 South Australia 和 Adelaide 的建筑市场需求和成本压力。我用 ABS 的 Building Approvals 作为未来建筑需求信号，用 Construction PPI 作为价格和成本压力信号。原始数据是 Excel，而且带有元数据行和宽表时间序列列，所以我用 Python 把它们清洗成长表。",
    )
    add_para(
        doc,
        "目标架构是 Databricks 加 Power BI。在 Databricks 中，数据可以先进入 Bronze 原始层，再清洗成 Silver 事实表，最后发布为 Gold fact 和 dimension tables。Power BI 里我建立了星型模型，包括 market indicators fact table、date dimension 和 series dimension。我也写了 DAX measures，例如 latest value、annual change、period change、building approvals 和 PPI index。",
    )
    add_para(
        doc,
        "报表一共有四页：管理层市场总览、建筑审批需求分析、建筑成本压力分析、数据质量与血缘。一个重要建模选择是使用 series_key，而不是直接使用 series_id，因为 ABS 的同一个 series_id 可能跨表重复。这样模型关系更可靠。整个项目展示了我能把业务理解、数据工程、SQL 建模、Power BI 报表和数据质量思维结合起来。",
    )

    add_heading(doc, "7. Short Answers for Interview Questions | 常见问题简答", 1)
    add_table(
        doc,
        ["Question | 问题", "Answer | 回答"],
        [
            ["Why use building approvals?", "They are an early demand signal for future construction activity. | 它是未来建筑活动和材料需求的前置信号。"],
            ["Why use PPI?", "It shows input and output price pressure in construction. | 它反映建筑投入和产出价格压力。"],
            ["Why split input and output PPI charts?", "They have different geographic grains: city-level and state-level. | 因为两者地理粒度不同，一个是城市，一个是州。"],
            ["Why use a star schema?", "It improves clarity, filtering, performance, and measure design. | 星型模型更清楚，也方便筛选、性能优化和写度量。"],
            ["Why use series_key?", "It prevents duplicate-key issues when the same series_id appears in multiple ABS tables. | 防止同一个 series_id 跨表重复导致关系不可靠。"],
            ["What would you improve next?", "Automate refresh in Databricks, add incremental loading, and publish the Power BI report to Service. | 下一步可做 Databricks 自动刷新、增量加载、发布到 Power BI Service。"],
        ],
        [2.1, 4.2],
    )

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_doc()
